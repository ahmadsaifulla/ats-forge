"""Resume optimization logic."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import re
from uuid import uuid4

from docx import Document

from app.models.schemas import OptimizedContentLine, OptimizedResume, OptimizedResumeSection
from app.utils.errors import OptimizationError
from app.utils.text import extract_bullets, extract_keywords, extract_sections, measurable_achievement_count

ACTION_VERBS = (
    "accelerated",
    "architected",
    "automated",
    "built",
    "collaborated",
    "created",
    "delivered",
    "designed",
    "developed",
    "drove",
    "enhanced",
    "executed",
    "implemented",
    "improved",
    "launched",
    "led",
    "managed",
    "optimized",
    "owned",
    "reduced",
    "scaled",
    "streamlined",
    "supported",
)


class ResumeOptimizerService:
    """Generate ATS-friendlier resume output without inventing experience."""

    def optimize(
        self,
        resume_id: str,
        resume_text: str,
        job_description: str,
        missing_keywords: list[str],
    ) -> tuple[OptimizedResume, bytes]:
        """Create a structured optimized resume and downloadable DOCX."""

        if not resume_text.strip():
            raise OptimizationError("Cannot optimize an empty resume.")

        sections = extract_sections(resume_text)
        optimized_sections = self._build_sections(sections, job_description, missing_keywords)
        plain_text = self._render_plain_text(optimized_sections)
        document_id = uuid4().hex
        docx_bytes = self._build_docx(optimized_sections)

        optimized_resume = OptimizedResume(
            resume_id=resume_id,
            document_id=document_id,
            plain_text=plain_text,
            sections=optimized_sections,
            disclaimer=(
                "This version is optimized for ATS readability using clearer structure, evidence-based phrasing, and only role terms "
                "that can be reasonably grounded in the source resume. Review every line and keep only claims that are accurate."
            ),
            generated_at=datetime.now(timezone.utc),
        )
        return optimized_resume, docx_bytes

    def _build_sections(
        self,
        sections: dict[str, str],
        job_description: str,
        missing_keywords: list[str],
    ) -> list[OptimizedResumeSection]:
        """Rewrite sections into a cleaner ATS-friendly format."""

        optimized: list[OptimizedResumeSection] = []
        role_keywords = extract_keywords(job_description, max_keywords=20)
        evidence_text = "\n".join(
            filter(
                None,
                [
                    sections.get("summary", ""),
                    sections.get("skills", ""),
                    sections.get("experience", ""),
                    sections.get("education", ""),
                    sections.get("other", ""),
                ],
            )
        )
        safe_keywords = self._select_grounded_keywords(role_keywords, evidence_text, missing_keywords)

        summary_lines = self._build_summary(
            sections=sections,
            role_keywords=role_keywords,
            safe_keywords=safe_keywords,
        )
        optimized.append(self._build_section("SUMMARY", summary_lines, label="Rewritten content"))

        skills_content = self._optimize_skills(sections.get("skills", ""), safe_keywords)
        if skills_content:
            optimized.append(self._build_section("SKILLS", skills_content, label="Rewritten content"))

        experience_content = self._optimize_experience(sections.get("experience", "") or sections.get("other", ""))
        if experience_content:
            optimized.append(self._build_section("EXPERIENCE", experience_content, label="Rewritten content"))

        education_content = self._normalize_paragraphs(sections.get("education", ""))
        if education_content:
            optimized.append(self._build_section("EDUCATION", education_content, label="Rewritten content"))

        suggested_additions = self._build_suggested_additions(missing_keywords)
        if suggested_additions:
            optimized.append(self._build_section("SUGGESTED ADDITIONS", suggested_additions, label="Suggested additions"))

        return optimized

    def _build_summary(
        self,
        *,
        sections: dict[str, str],
        role_keywords: list[str],
        safe_keywords: list[str],
    ) -> list[str]:
        """Create a concise, industry-standard summary based on resume evidence."""

        summary_lines: list[str] = []
        original_summary = self._normalize_paragraphs(sections.get("summary", ""))
        experience_text = sections.get("experience", "") or sections.get("other", "")
        measurable_count = measurable_achievement_count(experience_text)
        skills = self._optimize_skills(sections.get("skills", ""), safe_keywords)

        if original_summary:
            primary_line = self._trim_sentence(original_summary[0], limit=220)
            summary_lines.append(primary_line.rstrip(".") + ".")
        else:
            role_phrase = ", ".join(role_keywords[:3]) if role_keywords else "cross-functional delivery"
            summary_lines.append(
                f"Results-oriented professional with experience spanning {role_phrase} and ATS-friendly resume structure."
            )

        supporting_points: list[str] = []
        if measurable_count:
            supporting_points.append(
                f"Highlights {measurable_count} measurable achievement{'s' if measurable_count != 1 else ''} across prior experience."
            )
        if skills:
            supporting_points.append(
                "Core strengths include " + ", ".join(skills[:6]) + "."
            )
        if safe_keywords:
            supporting_points.append(
                "Role-aligned capabilities reflected in the resume include " + ", ".join(safe_keywords[:4]) + "."
            )

        return [summary_lines[0]] + supporting_points[:2]

    def _optimize_skills(self, skills_text: str, safe_keywords: list[str]) -> list[str]:
        """Build a clean skills list."""

        skill_items = []
        for line in skills_text.splitlines():
            fragments = re.split(r"[,|/]", line)
            skill_items.extend(part.strip() for part in fragments if part.strip())
        seen: set[str] = set()
        ordered: list[str] = []
        for item in skill_items + safe_keywords:
            cleaned = self._normalize_skill(item)
            lowered = cleaned.lower()
            if lowered not in seen:
                ordered.append(cleaned)
                seen.add(lowered)
        return ordered[:16]

    def _build_suggested_additions(self, missing_keywords: list[str]) -> list[str]:
        """Create honest gap suggestions without asserting unsupported experience."""

        return [
            f"Add {keyword} only if it is genuinely supported by your experience, projects, certifications, or training."
            for keyword in missing_keywords[:8]
        ]

    def _optimize_experience(self, experience_text: str) -> list[str]:
        """Rewrite bullets with clearer action-oriented phrasing."""

        bullets = extract_bullets(experience_text)
        if not bullets:
            paragraphs = self._normalize_paragraphs(experience_text)
            return [self._trim_sentence(line, limit=180) for line in paragraphs[:8]]

        rewritten: list[str] = []
        for bullet in bullets[:12]:
            cleaned = self._clean_bullet(bullet)
            if not cleaned:
                continue
            rewritten.append(cleaned)
        return rewritten

    def _normalize_paragraphs(self, text: str) -> list[str]:
        """Split freeform text into clean paragraph lines."""

        return [line.strip() for line in text.splitlines() if line.strip()]

    def _build_section(
        self,
        title: str,
        lines: list[str],
        *,
        label: str,
    ) -> OptimizedResumeSection:
        """Create a labeled section for the API response."""

        return OptimizedResumeSection(
            title=title,
            content=lines,
            labels=[label for _ in lines],
            items=[OptimizedContentLine(label=label, text=line) for line in lines],
        )

    def _select_grounded_keywords(
        self,
        role_keywords: list[str],
        evidence_text: str,
        missing_keywords: list[str],
    ) -> list[str]:
        """Keep only role keywords that are reasonably grounded in the source resume."""

        evidence_lower = evidence_text.lower()
        grounded: list[str] = []
        candidates = role_keywords + missing_keywords
        for keyword in candidates:
            normalized = keyword.lower().strip()
            if not normalized or normalized in grounded:
                continue
            parts = [part for part in re.split(r"\s+", normalized) if part]
            overlap = sum(1 for part in parts if part in evidence_lower)
            if normalized in evidence_lower or overlap >= max(1, len(parts) - 1):
                grounded.append(normalized)
        return grounded[:8]

    def _normalize_skill(self, skill: str) -> str:
        """Normalize skill text for cleaner competency lists."""

        cleaned = re.sub(r"\s+", " ", skill).strip(" .-")
        if not cleaned:
            return ""
        return cleaned[0].upper() + cleaned[1:]

    def _clean_bullet(self, bullet: str) -> str:
        """Rewrite a bullet to follow stronger resume conventions without changing claims."""

        cleaned = re.sub(r"\s+", " ", bullet).strip(" .-")
        cleaned = re.sub(r"^(responsible for|worked on|helped with|in charge of)\s+", "", cleaned, flags=re.I)
        cleaned = cleaned.strip()
        if not cleaned:
            return ""
        if cleaned[0].islower():
            cleaned = cleaned[0].upper() + cleaned[1:]
        if not any(cleaned.lower().startswith(verb) for verb in ACTION_VERBS):
            cleaned = self._promote_to_action_phrase(cleaned)
        return self._trim_sentence(cleaned.rstrip(".") + ".", limit=190)

    def _promote_to_action_phrase(self, bullet: str) -> str:
        """Add a stronger lead-in while preserving the original statement."""

        if re.match(r"^(improvement|increase|reduction|delivery|design|development)\b", bullet, flags=re.I):
            return f"Drove {bullet[0].lower() + bullet[1:]}"
        return f"Delivered {bullet[0].lower() + bullet[1:]}"

    def _trim_sentence(self, text: str, *, limit: int) -> str:
        """Trim overly long lines to a readable ATS-friendly length."""

        compact = re.sub(r"\s+", " ", text).strip()
        if len(compact) <= limit:
            return compact
        truncated = compact[: limit - 3].rsplit(" ", 1)[0]
        return truncated + "..."

    def _render_plain_text(self, sections: list[OptimizedResumeSection]) -> str:
        """Render optimized sections to plain text."""

        blocks = []
        for section in sections:
            body = "\n".join(f"- {item}" if len(section.content) > 1 else item for item in section.content)
            blocks.append(f"{section.title}\n{body}")
        return "\n\n".join(blocks).strip()

    def _build_docx(self, sections: list[OptimizedResumeSection]) -> bytes:
        """Build a DOCX document for download."""

        try:
            document = Document()
            for section in sections:
                document.add_heading(section.title, level=1)
                if section.title in {"SUMMARY", "EDUCATION"} and len(section.content) == 1:
                    document.add_paragraph(section.content[0])
                else:
                    for item in section.content:
                        document.add_paragraph(item, style="List Bullet")
            buffer = BytesIO()
            document.save(buffer)
        except Exception as exc:  # pragma: no cover - library exceptions vary
            raise OptimizationError("Unable to generate the optimized resume document.") from exc
        return buffer.getvalue()
